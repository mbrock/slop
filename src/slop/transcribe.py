import json
import logging
import re
import tempfile
from collections.abc import Awaitable, Callable
from contextlib import contextmanager, nullcontext
from io import BytesIO
from pathlib import Path
from subprocess import PIPE
from typing import Any, Literal, NotRequired, TypedDict

import anyio
from docx import Document
from docx.shared import Pt
from starlette.datastructures import UploadFile
from starlette.exceptions import HTTPException
from starlette.responses import RedirectResponse, Response, StreamingResponse
from tagflow import (
    attr,
    classes,
    tag,
    text,
)

from slop import app, gemini, rest
from slop.gemini import GeminiError, ModelOverloadedError
from slop.models import (
    Part,
    Tape,
    Utterance,
    get_tape,
    list_parts_for_tape,
    list_tapes,
    new_part_id,
    new_tape_id,
    save_part,
    save_tape,
)
from slop.parameter import Parameter
from slop.sndprompt import (
    improve_speaker_identification_part,
    transcribe_audio_part,
)
from slop.store import (
    ModelDecodeError,
    ModelNotFoundError,
    get_blob,
    save_blob,
)
from slop.views import speaker_classes, upload_area

logger = logging.getLogger("slop.transcribe")


MODEL_CHOICES = [
    "gemini-2.5-flash",
    "gemini-2.5-pro",
    "gemini-2.5-flash-lite",
]


TranscribePartCallable = Callable[
    [str, str, str, list[Part] | None],
    Awaitable[tuple[list[Utterance], str]],
]

transcribe_part_callable = Parameter[TranscribePartCallable](
    "app_transcribe_part"
)


class TapeJobPayload(TypedDict, total=False):
    kind: Literal["transcribe_next"]
    duration_seconds: int | str
    model_name: str


class PartJobPayload(TypedDict, total=False):
    kind: Literal["retranscribe", "improve_speakers", "update_speaker"]
    model_name: str
    hint: str
    utterance: int | str
    speaker: str


IMPROVE_SPEAKERS_ONCLICK = (
    "const hint = prompt('Enter any hints about the speakers (optional):'); "
    "if (hint === null) { return false; } "
    "const selected = document.querySelector(\"#model-selector input[name='model_name']:checked\"); "
    "const payload = { kind: 'improve_speakers', hint }; "
    "if (selected) { payload.model_name = selected.value; } "
    "this.setAttribute('hx-vals', JSON.stringify(payload)); "
    "return true;"
)


def load_tape_or_error(tape_id: str) -> Tape:
    """Return the tape or raise an appropriate HTTP error."""

    try:
        return get_tape(tape_id)
    except ModelNotFoundError as exc:
        raise HTTPException(status_code=404, detail="Tape not found") from exc
    except ModelDecodeError as exc:
        raise HTTPException(status_code=500, detail="Tape data invalid") from exc


def load_parts_or_error(tape: Tape) -> list[Part]:
    """Return ordered parts for ``tape`` or raise HTTP errors."""

    try:
        return list_parts_for_tape(tape)
    except ModelNotFoundError as exc:
        raise HTTPException(status_code=404, detail="Part not found") from exc


async def _read_payload() -> dict[str, Any]:
    """Normalise the incoming request payload to a plain dictionary."""

    req = app.request.get()
    content_type = req.headers.get("content-type", "")
    if "json" in content_type:
        data = await req.json()
        if not isinstance(data, dict):  # pragma: no cover - sanity check
            raise HTTPException(status_code=400, detail="JSON payload must be an object")
        return data

    form = await req.form()
    return {key: value for key, value in form.items()}


def _coerce_int(value: Any) -> int | None:
    if value is None:
        return None
    if isinstance(value, int):
        return value
    if isinstance(value, str):
        value = value.strip()
        if value.isdigit():
            return int(value)
    return None


@contextmanager
def layout(title: str):
    """Common layout wrapper for all pages."""
    with tag.html(lang="en"):
        with tag.head():
            with tag.title():
                text(f"{title} - Slop")

            # Tailwind + HTMX scripts
            with tag.script(src="https://cdn.tailwindcss.com"):
                pass
            with tag.script(src="https://unpkg.com/htmx.org@2.0.4"):
                pass
            with tag.script(src="https://unpkg.com/hyperscript.org@0.9.12"):
                pass

            # Add key tracking script
            with tag.script():
                text("""
                    window.keyPressed = undefined;
                    document.addEventListener('keydown', (e) => {
                        window.keyPressed = e.key;
                        console.log(window.keyPressed, "pressed");
                    });
                    document.addEventListener('keyup', () => {
                        window.keyPressed = undefined;
                        console.log("key released");
                    });
                """)

            # Add CSS for loading indicator
            with tag.style():
                text("""
                    .htmx-indicator {
                        opacity: 0;
                        transition: opacity 200ms ease-in;
                    }
                    .htmx-request .htmx-indicator {
                        opacity: 1
                    }
                    .htmx-request.htmx-indicator {
                        opacity: 1
                    }

                    button.htmx-request {
                        opacity: 0.5;
                        cursor: wait;
                    }
                """)

            with tag.script(
                type="module", src="https://cdn.jsdelivr.net/npm/media-chrome@3/+esm"
            ):
                pass

        with tag.body(classes="min-h-screen"):
            svg_icons()
            with tag.main():
                yield


def time_to_seconds(time_str: str) -> int:
    """Convert HH:MM:SS to total seconds."""
    h, m, s = map(int, time_str.split(":"))
    return h * 3600 + m * 60 + s


def calculate_progress(current: str, total: str) -> int:
    """Calculate progress percentage."""
    current_secs = time_to_seconds(current)
    total_secs = time_to_seconds(total)
    if total_secs == 0:
        return 0
    return int((current_secs / total_secs) * 100)


def render_tape_list() -> None:
    """Render the tape list content into the current TagFlow document."""
    with tag.div(classes="space-y-4", id="tape-list"):
        for tape in list_tapes():
            progress = calculate_progress(
                tape.current_position, tape.duration
            )
            with tag.div(classes="border rounded-lg p-4 hover:bg-gray-50"):
                with tag.a(
                    href=f"/tapes/{tape.id}",
                    classes="block",
                ):
                    with tag.div(classes="flex justify-between items-center mb-2"):
                        with tag.span(classes="font-medium"):
                            text(tape.filename)
                        with tag.span(classes="text-gray-500 text-sm"):
                            text(f"{tape.current_position} / {tape.duration}")

                    with tag.div(classes="bg-gray-200 rounded-full h-2"):
                        with tag.div(
                            classes="bg-blue-600 rounded-full h-2 transition-all",
                            style=f"width: {progress}%",
                        ):
                            pass


def render_home_content() -> None:
    """Render the shared home-page content."""
    breadcrumb({"Ieva's Tapes": "#"})
    with tag.div(classes="max-w-4xl mx-auto p-4"):
        with tag.div(classes="mb-8"):
            # with tag.h1(classes="text-2xl font-bold mb-4"):
            #     text("Ieva's Tapes")

            render_tape_list()

        with tag.div(classes="prose mx-auto"):
            upload_area(target="main")


def home():
    """Render the full home page layout with the upload area."""
    with layout("Home"):
        render_home_content()


def list_tapes_view(partial: bool = False) -> None:
    """Return either the full home view or just the list partial."""

    if partial:
        render_tape_list()
    else:
        home()


async def process_audio(input_path: Path) -> bytes:
    """
    Convert audio to Ogg Vorbis format using ffmpeg.
    Returns the processed audio as bytes.
    """
    logger.info(f"Processing audio file: {input_path}")
    process = await anyio.run_process(
        [
            "ffmpeg",
            "-i",
            str(input_path),  # Input file
            "-c:a",
            "libvorbis",  # Vorbis codec
            "-q:a",
            "8",  # Quality level (0-10, 4 is good)
            "-f",
            "ogg",  # Output format
            "pipe:1",  # Output to stdout
        ],
        stdout=PIPE,
        stderr=PIPE,
        check=False,
    )

    logger.info(f"FFmpeg process {process.returncode}")
    if process.returncode != 0:
        stderr = process.stderr.decode("utf-8", errors="replace")
        raise RuntimeError(f"FFmpeg failed: {stderr}")

    return process.stdout


async def _create_tape_from_upload(audio: UploadFile) -> Tape:
    """Persist an uploaded audio file and return the stored tape."""
    upload_name = audio.filename or "upload.ogg"

    with tempfile.NamedTemporaryFile(
        delete=False, suffix=Path(upload_name).suffix
    ) as tmp:
        content = await audio.read()
        tmp.write(content)
        tmp.flush()
        tmp_path = Path(tmp.name)

        try:
            # Get audio duration
            duration = await get_audio_duration(tmp_path)

            # Process and store the audio
            processed_audio = await process_audio(tmp_path)
            audio_hash = save_blob(processed_audio, "audio/ogg")

            # Create tape record
            tape_id = new_tape_id()
            tape = Tape(
                id=tape_id,
                filename=upload_name,
                audio_hash=audio_hash,
                duration=duration,
            )
            save_tape(tape)
        finally:
            tmp_path.unlink()

    return tape


async def create_tape(audio: UploadFile) -> Response:
    """Handle a tape upload and respond with its identifier."""

    tape = await _create_tape_from_upload(audio)
    payload = json.dumps({"id": tape.id})
    return Response(
        content=payload,
        media_type="application/json",
        status_code=201,
        headers={"HX-Redirect": f"/tapes/{tape.id}"},
    )


def svg_icons():
    with tag("svg", classes="hidden"):
        # backward symbol
        with tag(
            "symbol",
            id="backward",
            viewBox="0 0 24 24",
            **{
                "stroke-width": "1.5",
                "stroke-linecap": "round",
                "stroke-linejoin": "round",
            },
        ):
            with tag(
                "path",
                d="M8 5L5 8M5 8L8 11M5 8H13.5C16.5376 8 19 10.4624 19 13.5C19 15.4826 18.148 17.2202 17 18.188",
            ):
                pass
            with tag("path", d="M5 15V19"):
                pass
            with tag(
                "path",
                d="M8 18V16C8 15.4477 8.44772 15 9 15H10C10.5523 15 11 15.4477 11 16V18C11 18.5523 10.5523 19 10 19H9C8.44772 19 8 18.5523 8 18Z",
            ):
                pass

        # play symbol
        with tag("symbol", id="play", viewBox="0 0 24 24"):
            with tag(
                "path",
                **{
                    "fill-rule": "evenodd",
                    "clip-rule": "evenodd",
                    "d": "M4.5 5.653c0-1.426 1.529-2.33 2.779-1.643l11.54 6.348c1.295.712 1.295 2.573 0\
        3.285L7.28 19.991c-1.25.687-2.779-.217-2.779-1.643V5.653z",
                },
            ):
                pass

        # pause symbol
        with tag("symbol", id="pause", viewBox="0 0 24 24"):
            with tag(
                "path",
                **{
                    "fill-rule": "evenodd",
                    "clip-rule": "evenodd",
                    "d": "M6.75 5.25a.75.75 0 01.75-.75H9a.75.75 0 01.75.75v13.5a.75.75 0\
      01-.75.75H7.5a.75.75 0 01-.75-.75V5.25zm7.5 0A.75.75 0 0115 4.5h1.5a.75.75 0 01.75.75v13.5a.75.75 0\
      01-.75.75H15a.75.75 0 01-.75-.75V5.25z",
                },
            ):
                pass

        # forward symbol
        with tag(
            "symbol",
            id="forward",
            viewBox="0 0 24 24",
            **{
                "stroke-width": "1.5",
                "stroke-linecap": "round",
                "stroke-linejoin": "round",
            },
        ):
            with tag(
                "path",
                d="M16 5L19 8M19 8L16 11M19 8H10.5C7.46243 8 5 10.4624 5 13.5C5 15.4826 5.85204 17.2202 7 18.188",
            ):
                pass
            with tag("path", d="M13 15V19"):
                pass
            with tag(
                "path",
                d="M16 18V16C16 15.4477 16.4477 15 17 15H18C18.5523 15 19 15.4477 19 16V18C19 18.5523 18.5523 19 18\
      19H17C16.4477 19 16 18.5523 16 18Z",
            ):
                pass

        # high symbol
        with tag("symbol", id="high", viewBox="0 0 24 24"):
            with tag(
                "path",
                d="M13.5 4.06c0-1.336-1.616-2.005-2.56-1.06l-4.5 4.5H4.508c-1.141 0-2.318.664-2.66 1.905A9.76 9.76 0\
      001.5 12c0 .898.121 1.768.35 2.595.341 1.24 1.518 1.905 2.659 1.905h1.93l4.5 4.5c.945.945 2.561.276\
      2.561-1.06V4.06zM18.584 5.106a.75.75 0 011.06 0c3.808 3.807 3.808 9.98 0 13.788a.75.75 0 11-1.06-1.06\
      8.25 8.25 0 000-11.668.75.75 0 010-1.06z",
            ):
                pass
            with tag(
                "path",
                d="M15.932 7.757a.75.75 0 011.061 0 6 6 0 010 8.486.75.75 0 01-1.06-1.061 4.5 4.5 0 000-6.364.75.75 0\
      010-1.06z",
            ):
                pass

        # off symbol
        with tag("symbol", id="off", viewBox="0 0 24 24"):
            with tag(
                "path",
                d="M13.5 4.06c0-1.336-1.616-2.005-2.56-1.06l-4.5 4.5H4.508c-1.141 0-2.318.664-2.66 1.905A9.76 9.76 0\
      001.5 12c0 .898.121 1.768.35 2.595.341 1.24 1.518 1.905 2.659 1.905h1.93l4.5 4.5c.945.945 2.561.276\
      2.561-1.06V4.06zM17.78 9.22a.75.75 0 10-1.06 1.06L18.44 12l-1.72 1.72a.75.75 0 001.06 1.06l1.72-1.72 1.72\
      1.72a.75.75 0 101.06-1.06L20.56 12l1.72-1.72a.75.75 0 00-1.06-1.06l-1.72 1.72-1.72-1.72z",
            ):
                pass


def audio_player(src: str):
    with tag("media-controller", audio=True, classes="w-full"):
        attr(
            "style",
            """
            --media-background-color: transparent;
            --media-control-background: transparent;
            --media-control-hover-background: transparent;
        """,
        )
        with tag("audio", slot="media", src=src, crossorigin=True):
            pass
        with tag(
            "media-control-bar",
            classes="h-12 w-full bg-white items-center",
        ):
            # rounded-md ring-1 ring-slate-700/10 shadow-xl shadow-black/5
            with tag("media-seek-backward-button", classes="p-0"):
                with tag.svg(
                    slot="icon",
                    aria_hidden=True,
                    classes="w-7 h-7 fill-none stroke-gray-500",
                ):
                    with tag.use(href="#backward"):
                        pass
            with tag(
                "media-play-button",
                classes="h-7 w-7 p-2 mx-3 rounded-full bg-gray-700",
            ):
                with tag.svg(slot="play", aria_hidden=True, classes="relative left-px"):
                    with tag.use(href="#play"):
                        pass
                with tag.svg(slot="pause", aria_hidden=True):
                    with tag.use(href="#pause"):
                        pass
            with tag("media-seek-forward-button", classes="p-0"):
                with tag.svg(
                    slot="icon",
                    aria_hidden=True,
                    classes="w-7 h-7 fill-none stroke-gray-500",
                ):
                    with tag.use(href="#forward"):
                        pass
            with tag("media-time-display", classes="text-gray-500 text-sm"):
                pass
            with tag(
                "media-time-range",
                classes="block h-2 min-h-0 p-0 m-2 rounded-md bg-gray-50",
            ):
                attr(
                    "style",
                    """
                    --media-range-track-background: transparent;
                    --media-time-range-buffered-color: rgb(0 0 0 / 0.02);
                    --media-range-bar-color: rgb(79 70 229);
                    --media-range-track-border-radius: 4px;
                    --media-range-track-height: 0.5rem;
                    --media-range-thumb-background: rgb(79 70 229);
                    --media-range-thumb-box-shadow: 0 0 0 2px rgb(255 255 255 / 0.9);
                    --media-range-thumb-width: 0.25rem;
                    --media-range-thumb-height: 1rem;
                    --media-preview-time-text-shadow: transparent;
                """,
                )
                with tag("media-preview-time-display", classes="text-gray-600 text-xs"):
                    pass
            with tag("media-duration-display", classes="text-gray-500 text-xs"):
                pass
            with tag("media-mute-button"):
                with tag.svg(
                    slot="high",
                    aria_hidden=True,
                    classes="w-5 h-5 fill-gray-500",
                ):
                    with tag.use(href="#high"):
                        pass
                with tag.svg(
                    slot="medium",
                    aria_hidden=True,
                    classes="w-5 h-5 fill-gray-500",
                ):
                    with tag.use(href="#high"):
                        pass
                with tag.svg(
                    slot="low",
                    aria_hidden=True,
                    classes="w-5 h-5 fill-gray-500",
                ):
                    with tag.use(href="#high"):
                        pass
                with tag.svg(
                    slot="off",
                    aria_hidden=True,
                    classes="w-5 h-5 fill-gray-500",
                ):
                    with tag.use(href="#off"):
                        pass


def breadcrumb(items: dict[str, str]):
    with tag.nav(
        classes="hidden sm:flex bg-gray-300 px-4 py-1 border-b border-gray-400",
        aria_label="Breadcrumb",
    ):
        with tag.ol(role="list", classes="flex items-center space-x-4"):
            total = len(items)
            for i, (label, href) in enumerate(items.items()):
                with tag.li():
                    with tag.div(classes="flex items-center"):
                        # Render the arrow separator only if it's not the first item
                        if i > 0:
                            with tag.svg(
                                classes="size-5 shrink-0 text-gray-400 mr-4",
                                viewBox="0 0 20 20",
                                fill="currentColor",
                                aria_hidden="true",
                                **{"data-slot": "icon"},
                            ):
                                with tag.path(
                                    **{
                                        "fill-rule": "evenodd",
                                        "d": "M8.22 5.22a.75.75 0 0 1 1.06 0l4.25 4.25a.75.75 0 0 1 0 1.06l-4.25 4.25a.75.75 0 0 1-1.06-1.06L11.94 10 8.22 6.28a.75.75 0 0 1 0-1.06Z",
                                        "clip-rule": "evenodd",
                                    }
                                ):
                                    pass
                        with tag.a(
                            href=href,
                            classes="text-sm font-medium text-gray-600 hover:text-gray-700",
                            aria_current="page" if i == (total - 1) else None,
                        ):
                            text(label)


def button_view(label: str, href: str | None = None, type: str = "button", **attrs):
    """Renders a button with consistent styling."""
    base_classes = "inline-flex items-center rounded-md px-3 py-1 text-sm font-semibold text-gray-900 bg-white shadow-sm ring-1 ring-inset ring-gray-300 hover:bg-gray-50 whitespace-nowrap"

    if href:
        with tag.a(href=href, classes=base_classes):
            text(label)
    else:
        with tag.button(type=type, classes=base_classes, **attrs):
            text(label)


def tape_header(title: str, tape_id: str):
    """Renders the tape header with title and action buttons."""
    tape = load_tape_or_error(tape_id)
    context_parts_value = str(tape.context_parts)

    with tag.div():
        breadcrumb({"Ieva's Tapes": "/", title: "#"})
        with tag.div(
            classes="mt-2 md:flex md:items-center md:justify-between px-4 py-2"
        ):
            with tag.div(classes="min-w-0 flex-1"):
                with tag.h2(
                    classes="text-2xl/7 font-bold text-gray-900 sm:truncate sm:text-3xl sm:tracking-tight"
                ):
                    text(title)
            with tag.div(classes="mt-4 flex shrink-0 gap-4 md:ml-4 md:mt-0"):
                # Context parts control
                with tag.form(
                    classes="flex items-center gap-2",
                    **{"hx-patch": f"/tapes/{tape_id}", "hx-swap": "none"},
                ):
                    with tag.label(classes="text-sm text-gray-600"):
                        text("Context parts:")
                    with tag.input(
                        type="number",
                        name="context_parts",
                        min="0",
                        max="5",
                        value=context_parts_value,
                        classes="w-16 px-2 py-1 text-sm border rounded",
                        **{
                            "hx-trigger": "change",
                            "hx-vals": "js:{context_parts: parseInt(this.value, 10)}",
                        },
                    ):
                        pass

                # Model selector
                with tag.form(
                    id="model-selector",
                    classes="flex items-center gap-2",
                ):
                    with tag.label(classes="text-sm text-gray-600"):
                        text("Model:")
                    with tag.div(classes="flex gap-2"):
                        for model in MODEL_CHOICES:
                            with tag.label(classes="flex items-center gap-1"):
                                with tag.input(
                                    type="radio",
                                    name="model_name",
                                    value=model,
                                    checked=(model == gemini.model.get()),
                                    classes="text-blue-600",
                                ):
                                    pass
                                with tag.span(classes="text-sm text-gray-600"):
                                    text(model)

                with tag.form(
                    onsubmit="const name = prompt('New name:', this.querySelector('input').value); if (!name) return false; this.querySelector('input').value = name;",
                    **{"hx-patch": f"/tapes/{tape_id}", "hx-swap": "none"},
                ):
                    with tag.input(type="hidden", name="name", value=title):
                        pass
                    button_view("Rename", type="submit")
                button_view("Export DOCX", href=f"/tapes/{tape_id}/export")


def render_part(tape_id: str, part_index: int, part: Part) -> None:
    part_path = f"/parts/{part.id}"

    with tag.div(
        id=f"part-{part_index}",
        classes="flex flex-col gap-2 p-4 py-2 border-t-4 border-gray-400 mt-4",
    ):
        with tag.div(classes="flex items-center gap-2"):
            if part.audio_hash:
                audio_player(f"/audio/{part.audio_hash}")

            button_view(
                "Edit",
                **{
                    "hx-get": f"{part_path}/edit",
                    "hx-target": f"#part-content-{part_index}",
                    "hx-swap": "innerHTML",
                },
            )

            button_view(
                "Retranscribe",
                **{
                    "hx-post": f"{part_path}/jobs",
                    "hx-target": f"#part-{part_index}",
                    "hx-swap": "outerHTML",
                    "hx-include": "#model-selector",
                    "hx-vals": '{"kind": "retranscribe"}',
                },
            )

            button_view(
                "Improve Speakers",
                **{
                    "hx-post": f"{part_path}/jobs",
                    "hx-target": f"#part-{part_index}",
                    "hx-swap": "outerHTML",
                    "hx-include": "#model-selector",
                    "onclick": IMPROVE_SPEAKERS_ONCLICK,
                },
            )

        with tag.div(id=f"part-content-{part_index}", classes="w-full"):
            with tag.div(classes="flex flex-wrap gap-4"):
                for i, utterance in enumerate(part.utterances):
                    render_utterance(part, part_index, i, utterance)


def render_parts_container() -> None:
    tape = app.tape.get()
    tape_id = tape.id
    parts = load_parts_or_error(tape)

    with tag.div(id="parts", classes="flex flex-col gap-2"):
        if parts:
            for i, part in enumerate(parts):
                render_part(tape_id, i, part)


def list_parts_view() -> None:
    """HTMX partial rendering the full parts container."""

    render_parts_container()


def view_part():
    """Renders a single part as a partial view."""
    tape = app.tape.get()
    tape_id = tape.id
    part_index = app.part_index.get()
    part = app.get_part()

    render_part(tape_id, part_index, part)


def render_utterance(part: Part, part_index: int, i: int, utterance):
    with tag.span(
        **{
            "data-speaker": utterance.speaker,
            "hx-post": f"/parts/{part.id}/jobs",
            "hx-trigger": "click",
            "hx-vals": f'js:{{"kind": "update_speaker", "utterance": {i}, "speaker": window.keyPressed}}',
            "hx-swap": "outerHTML",
            "onclick": "if (!window.keyPressed) { return false; }",
        }
    ):
        classes(speaker_classes(utterance.speaker), "hover:underline")
        text(utterance.text)


def edit_part_dialog():
    """Renders the inline edit form for a part."""
    tape = app.tape.get()
    part = app.get_part()
    part_index = app.part_index.get()

    # Convert utterances to text format
    text_content = "\n\n".join(f"{u.speaker}: {u.text}" for u in part.utterances)

    with tag.form(
        classes="space-y-4",
        **{
            "hx-patch": f"/parts/{part.id}",
            "hx-target": f"#part-content-{part_index}",
            "hx-swap": "innerHTML",
        },
    ):
        with tag.div(classes="flex gap-2"):
            with tag.textarea(
                name="content",
                classes="flex-1 h-[60vh] font-mono p-2 border rounded text-sm",
                placeholder="Format: 'S1: Hello\n\nS2: Hi there'",
            ):
                text(text_content)
            with tag.div(classes="flex flex-col gap-2"):
                with tag.button(
                    type="submit",
                    classes="p-2 text-blue-600 hover:text-blue-800",
                ):
                    with tag.svg(
                        xmlns="http://www.w3.org/2000/svg",
                        fill="none",
                        viewBox="0 0 24 24",
                        **{"stroke-width": "1.5"},
                        stroke="currentColor",
                        classes="w-5 h-5",
                    ):
                        with tag.path(
                            **{
                                "stroke-linecap": "round",
                                "stroke-linejoin": "round",
                                "d": "m4.5 12.75 6 6 9-13.5",
                            }
                        ):
                            pass

                with tag.button(
                    type="button",
                    classes="p-2 text-gray-600 hover:text-gray-800",
                    **{
                        "hx-get": f"/parts/{part.id}",
                        "hx-target": f"#part-{part_index}",
                        "hx-swap": "outerHTML",
                    },
                ):
                    with tag.svg(
                        xmlns="http://www.w3.org/2000/svg",
                        fill="none",
                        viewBox="0 0 24 24",
                        **{"stroke-width": "1.5"},
                        stroke="currentColor",
                        classes="w-5 h-5",
                    ):
                        with tag.path(
                            **{
                                "stroke-linecap": "round",
                                "stroke-linejoin": "round",
                                "d": "M6 18 18 6M6 6l12 12",
                            }
                        ):
                            pass


def view_tape():
    """
    Renders the tape page, showing the audio player and parts.
    """
    tape = app.tape.get()
    tape_id = tape.id
    default_part_length = str(app.part_duration_seconds.get())

    with layout(f"Tape - {tape.filename}"):
        with tag.div(classes="prose mx-auto"):
            tape_header(tape.filename, tape.id)

            render_parts_container()

            # Progress bar and transcribe button section
            with tag.div(classes="border-t-4 border-gray-400 mt-4"):
                with tag.div(classes="flex items-center justify-between p-4"):
                    # Progress info
                    with tag.div(classes="flex items-center gap-2"):
                        with tag.span(classes="text-gray-500 text-sm"):
                            text(f"{tape.current_position} / {tape.duration}")
                        with tag.div(classes="w-48 bg-gray-200 rounded-full h-2"):
                            progress = calculate_progress(
                                tape.current_position, tape.duration
                            )
                            with tag.div(
                                classes="bg-blue-600 rounded-full h-2 transition-all",
                                style=f"width: {progress}%",
                            ):
                                pass

                    with tag.div(classes="flex items-center gap-3"):
                        with tag.form(
                            id="transcribe-options",
                            classes="flex items-center gap-2",
                        ):
                            with tag.label(classes="text-sm text-gray-600"):
                                text("Part length (s):")
                            with tag.input(
                                type="number",
                                name="duration_seconds",
                                min="5",
                                max="600",
                                step="5",
                                value=default_part_length,
                                classes="w-20 px-2 py-1 text-sm border rounded",
                            ):
                                pass

                        # Transcribe button
                        button_view(
                            "Transcribe more",
                            **{
                                "hx-post": f"/tapes/{tape_id}/jobs",
                                "hx-target": "#parts",
                                "hx-swap": "beforeend",
                                "hx-include": "#model-selector, #transcribe-options",
                                "hx-vals": '{"kind": "transcribe_next"}',
                            },
                        )


def _serve_audio_hash(hash_: str) -> Response:
    req = app.request.get()
    range_header = req.headers.get("range")

    try:
        data, mime_type = get_blob(hash_)
    except KeyError:
        raise HTTPException(status_code=404, detail="Audio not found")
    file_size = len(data)

    if not range_header:
        return Response(
            content=data,
            media_type=mime_type,
            headers={"accept-ranges": "bytes", "content-length": str(file_size)},
        )

    try:
        range_str = range_header.replace("bytes=", "")
        start_str, end_str = range_str.split("-")
        start = int(start_str) if start_str else 0
        end = int(end_str) if end_str else file_size - 1
    except ValueError:
        raise HTTPException(status_code=416, detail="Invalid range header")

    if start >= file_size or end >= file_size or start > end:
        raise HTTPException(
            status_code=416, detail=f"Range not satisfiable. File size: {file_size}"
        )

    content_length = end - start + 1

    return Response(
        content=data[start : end + 1],
        status_code=206,
        media_type=mime_type,
        headers={
            "content-range": f"bytes {start}-{end}/{file_size}",
            "accept-ranges": "bytes",
            "content-length": str(content_length),
        },
    )


def get_audio():
    """Serve audio file by hash with support for range requests."""

    hash_ = app.request.get().path_params["hash_"]
    return _serve_audio_hash(hash_)


def get_tape_media() -> Response:
    """Return the canonical audio for the current tape."""

    tape = app.tape.get()
    if not tape.audio_hash:
        raise HTTPException(status_code=404, detail="Tape has no audio")

    return _serve_audio_hash(tape.audio_hash)


async def transcribe_next_part(
    duration_seconds: int | None = None,
    model_name: str | None = None,
):
    """
    Transcribe the next audio part for the given tape.

    Optional form fields allow callers to override the part duration (in
    seconds) and select a Gemini model for this request.
    """
    tape = app.tape.get()
    tape_id = tape.id

    # Get context parts
    existing_parts = load_parts_or_error(tape)
    context_parts: list[Part] = []
    if existing_parts:
        start_idx = max(0, len(existing_parts) - tape.context_parts)
        context_parts = existing_parts[start_idx:]

    # Calculate part times
    start_time = tape.current_position
    base_duration = app.part_duration_seconds.get()
    requested_seconds = (
        duration_seconds if duration_seconds is not None else base_duration
    )
    part_seconds = max(1, min(requested_seconds, 15 * 60))
    end_time = increment_time(start_time, seconds=part_seconds)

    # Optionally switch the Gemini model just for this request
    chosen_model = model_name if model_name in MODEL_CHOICES else None
    model_context = (
        gemini.model.using(chosen_model)
        if chosen_model and chosen_model != gemini.model.get()
        else nullcontext()
    )

    try:
        with model_context:
            transcribe_fn = (
                transcribe_part_callable.peek() or transcribe_audio_part
            )
            # Transcribe the part
            utterances, part_hash = await transcribe_fn(
                tape_id,
                start_time,
                end_time,
                context_parts,
            )

        # Create and save the new part
        part = Part(
            id=new_part_id(),
            tape_id=tape_id,
            start_time=start_time,
            end_time=end_time,
            audio_hash=part_hash,
            utterances=utterances,
        )
        save_part(part)

        # Advance current position while leaving a 1-second overlap for continuity
        advance_seconds = max(part_seconds - 1, 0)
        tape.current_position = increment_time(start_time, seconds=advance_seconds)

        # Append part ID to tape and save
        tape.part_ids.append(part.id)
        save_tape(tape)

        # Return the new part as a partial view
        render_part(tape_id, len(tape.part_ids) - 1, part)

    except ModelOverloadedError as e:
        with tag.div(
            classes="fixed bottom-4 right-4 bg-yellow-100 border border-yellow-400 text-yellow-700 px-4 py-3 rounded shadow-lg",
            role="alert",
            **{
                "hx-swap-oob": "true",
                "_": "on load wait 5s then remove me",
            },
        ):
            with tag.p(classes="font-bold"):
                text("Model overloaded")
            with tag.p():
                text(
                    f"{gemini.model.get()} is busy. Try switching to {e.alternative_model} and retry."
                )

    except GeminiError as e:
        # Show error toast notification
        with tag.div(
            classes="fixed bottom-4 right-4 bg-red-100 border border-red-400 text-red-700 px-4 py-3 rounded shadow-lg",
            role="alert",
            **{
                "hx-swap-oob": "true",
                "_": "on load wait 5s then remove me",
            },
        ):
            with tag.p(classes="font-bold"):
                text("Error")
            with tag.p():
                text(e.message)


async def retranscribe_part(model_name: str | None = None):
    """Retranscribe a specific part using Gemini."""
    tape = app.tape.get()
    tape_id = tape.id
    part_index = app.part_index.get()
    part = app.get_part()

    # Get context parts
    all_parts = load_parts_or_error(tape)
    context_parts: list[Part] = []
    if part_index > 0 and all_parts:
        start_idx = max(0, part_index - tape.context_parts)
        context_parts = all_parts[start_idx:part_index]

    model_context = (
        gemini.model.using(model_name)
        if model_name in MODEL_CHOICES and model_name != gemini.model.get()
        else nullcontext()
    )

    with model_context:
        utterances, part_hash = await transcribe_audio_part(
            tape_id,
            part.start_time,
            part.end_time,
            context_parts,
        )

    # Update the part
    part.audio_hash = part_hash
    part.utterances = utterances
    save_part(part)

    # Return the updated part view
    render_part(tape_id, part_index, part)


async def improve_part_speakers(hint: str | None = None):
    """Improve speaker identification for a specific part using Gemini."""
    tape = app.tape.get()
    tape_id = tape.id
    part_index = app.part_index.get()
    part = app.get_part()

    # Get context parts
    all_parts = load_parts_or_error(tape)
    context_parts: list[Part] = []
    if part_index > 0 and all_parts:
        start_idx = max(0, part_index - tape.context_parts)
        context_parts = all_parts[start_idx:part_index]

    # Improve speaker identification
    utterances = await improve_speaker_identification_part(
        tape,
        part,
        part.utterances,
        context_parts,
        hint,
    )

    # Update the part
    part.utterances = utterances
    save_part(part)

    # Return the updated part view
    render_part(tape_id, part_index, part)


def increment_time(time_str: str, seconds: int) -> str:
    """Add `seconds` to a time string in HH:MM:SS format."""
    h, m, s_ = map(int, time_str.split(":"))
    total_seconds = h * 3600 + m * 60 + s_ + seconds
    h, remainder = divmod(total_seconds, 3600)
    m, s_ = divmod(remainder, 60)
    return f"{h:02d}:{m:02d}:{s_:02d}"


async def get_audio_duration(input_path: Path) -> str:
    """
    Get the duration of an audio file using ffprobe.
    Returns duration in HH:MM:SS format.
    """
    process = await anyio.run_process(
        [
            "ffprobe",
            "-v",
            "quiet",
            "-show_entries",
            "format=duration",
            "-of",
            "default=noprint_wrappers=1:nokey=1",
            str(input_path),
        ],
        stdout=PIPE,
        stderr=PIPE,
        check=False,
    )

    if process.returncode != 0:
        stderr = process.stderr.decode("utf-8", errors="replace")
        raise RuntimeError(f"FFprobe failed: {stderr}")

    duration_secs = float(process.stdout.decode().strip())
    hours = int(duration_secs // 3600)
    minutes = int((duration_secs % 3600) // 60)
    seconds = int(duration_secs % 60)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"


def update_part(content: str):
    """Updates a part's utterances from the edit form."""
    tape = app.tape.get()
    part_index = app.part_index.get()
    part = app.get_part()

    # Get the last speaker from the previous part if it exists
    last_speaker = None
    if part_index > 0:
        all_parts = load_parts_or_error(tape)
        if part_index - 1 < len(all_parts):
            previous = all_parts[part_index - 1]
            if previous.utterances:
                last_speaker = previous.utterances[-1].speaker

    # Parse the content into utterances
    utterances = []
    current_speaker = last_speaker or "S1"  # Default to S1 if no previous speaker
    speaker_pattern = re.compile(r"^\s*(\w+)\s*:\s*(.*)$")

    for line in content.strip().split("\n"):
        line = line.strip()
        if not line:
            continue

        if match := speaker_pattern.match(line):
            # Line starts with a speaker tag
            current_speaker = match.group(1)
            text = match.group(2).strip()
        else:
            # Line continues with current speaker
            text = line

        if text:  # Only add if there's actual text content
            utterances.append(
                Utterance(
                    speaker=current_speaker,
                    text=text,
                )
            )

    # Update the part
    part.utterances = utterances
    save_part(part)

    # Return the updated part view
    with tag.div(classes="flex flex-wrap gap-4"):
        for i, utterance in enumerate(part.utterances):
            render_utterance(part, part_index, i, utterance)


def rename_tape(new_name: str) -> None:
    """Update the tape's display name."""
    tape = app.tape.get()
    if not new_name:
        raise HTTPException(status_code=400, detail="New name required")

    tape.filename = new_name
    save_tape(tape)


def export_tape():
    """Export the tape as a DOCX file."""
    tape = app.tape.get()

    # Create a new document
    doc = Document()

    # Add title
    title = doc.add_heading(tape.filename, level=1)
    title.runs[0].font.size = Pt(16)

    # Add each part
    for part in load_parts_or_error(tape):
        # Add timestamp as a subheading
        doc.add_heading(f"{part.start_time} - {part.end_time}", level=2).runs[
            0
        ].font.size = Pt(12)

        # Add utterances
        for utterance in part.utterances:
            p = doc.add_paragraph()
            speaker_run = p.add_run(f"{utterance.speaker}: ")
            speaker_run.bold = True if utterance.speaker == "S1" else False
            text_run = p.add_run(utterance.text)
            text_run.bold = True if utterance.speaker == "S1" else False

    # Save to BytesIO
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)

    # Return as downloadable file
    # Normalize filename: replace spaces with underscores and remove/replace any problematic characters
    safe_filename = re.sub(r"[^\w\-\.]", "_", tape.filename.replace(" ", "_"))
    # Ensure ASCII encoding for Content-Disposition header
    encoded_filename = safe_filename.encode("ascii", "ignore").decode("ascii")
    filename = f"{encoded_filename}.docx"

    return StreamingResponse(
        docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def update_context_parts(context_parts: int) -> None:
    """Updates the number of context parts to use for transcription."""
    tape = app.tape.get()

    tape.context_parts = max(
        0, min(5, context_parts)
    )  # Clamp between 0 and 5
    save_tape(tape)


async def patch_tape() -> Response:
    """Handle a partial update to the tape resource."""

    payload = await _read_payload()
    updated = False

    if (name := payload.get("name")) is not None:
        rename_tape(str(name))
        updated = True

    if "context_parts" in payload:
        context_value = _coerce_int(payload.get("context_parts"))
        if context_value is None:
            raise HTTPException(
                status_code=400, detail="context_parts must be an integer"
            )
        update_context_parts(context_value)
        updated = True

    if not updated:
        raise HTTPException(status_code=400, detail="No valid fields to update")

    return Response(status_code=204)


def update_speaker(utterance_index: int, speaker: str):
    """Updates the speaker for a specific utterance."""
    part_index = app.part_index.get()
    part = app.get_part()

    try:
        utterance = part.utterances[utterance_index]
    except IndexError:
        raise HTTPException(status_code=404, detail="Utterance not found")

    speaker_value = speaker.strip()
    if not speaker_value:
        raise HTTPException(status_code=400, detail="Speaker is required")

    if speaker_value.isdigit():
        speaker_value = f"S{speaker_value}"

    utterance.speaker = speaker_value
    save_part(part)

    # Return the updated utterance view
    render_utterance(part, part_index, utterance_index, utterance)


async def tape_jobs() -> Response | None:
    """Fan-in for tape-level job requests."""

    payload = await _read_payload()
    kind = payload.get("kind")

    if kind == "transcribe_next":
        duration = _coerce_int(payload.get("duration_seconds"))
        model_name = payload.get("model_name")
        return await transcribe_next_part(
            duration_seconds=duration,
            model_name=model_name if isinstance(model_name, str) else None,
        )

    raise HTTPException(status_code=400, detail="Unsupported job kind")


async def part_jobs() -> Response | None:
    """Handle part-specific job invocations."""

    payload = await _read_payload()
    kind = payload.get("kind")

    if kind == "retranscribe":
        model_name = payload.get("model_name")
        return await retranscribe_part(
            model_name=model_name if isinstance(model_name, str) else None
        )

    if kind == "improve_speakers":
        hint = payload.get("hint")
        model_name = payload.get("model_name")
        return await improve_part_speakers(
            hint=hint if isinstance(hint, str) else None,
            model_name=model_name if isinstance(model_name, str) else None,
        )

    if kind == "update_speaker":
        utterance = _coerce_int(payload.get("utterance"))
        speaker = payload.get("speaker")
        if utterance is None or not isinstance(speaker, str):
            raise HTTPException(
                status_code=400,
                detail="update_speaker requires 'utterance' and 'speaker'",
            )
        update_speaker(utterance, speaker)
        return None

    raise HTTPException(status_code=400, detail="Unsupported job kind")
